"""
Exportação do relatório de conformidade.

Formatos: Markdown (``.md``), Excel (``.xlsx``), Word (``.docx``) e PDF.
Alinhado ao frontend: paleta Codevasf, cards de resumo, legendas e status.
"""

from __future__ import annotations

from datetime import datetime
from io import BytesIO
from pathlib import Path

from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter

from conformidade.checklist import label_tipo
from conformidade.models import RelatorioConformidade, StatusConformidade

# Paleta institucional (manual Codevasf + frontend)
AZUL_MARCA = "005CA8"
VERDE_MARCA = "007D4E"
AZUL_ESCURO = "222B54"
OFFWHITE = "F2F2F2"
BORDA = "B7D7E8"

STATUS_LABEL = {
    StatusConformidade.ATENDIDO: "ATENDIDO",
    StatusConformidade.PARCIAL: "PARCIAL",
    StatusConformidade.NAO_ATENDIDO: "NÃO ATENDIDO",
}

STATUS_LEGENDA = {
    StatusConformidade.ATENDIDO: "Documento Ok",
    StatusConformidade.PARCIAL: "Incompleto / Dúvida",
    StatusConformidade.NAO_ATENDIDO: "Ausente",
}

# Fundos dos badges de status (como .cv-badge-* no frontend)
STATUS_FILL = {
    StatusConformidade.ATENDIDO: PatternFill("solid", fgColor="D7F0E5"),
    StatusConformidade.PARCIAL: PatternFill("solid", fgColor="EEF7D4"),
    StatusConformidade.NAO_ATENDIDO: PatternFill("solid", fgColor="FDE2E2"),
}

# Fundos dos cards de resumo (como .cv-card-stat no frontend)
CARD_FILL = {
    StatusConformidade.ATENDIDO: PatternFill("solid", fgColor="EAF8F0"),
    StatusConformidade.PARCIAL: PatternFill("solid", fgColor="F7F9E8"),
    StatusConformidade.NAO_ATENDIDO: PatternFill("solid", fgColor="FDECED"),
}

STATUS_RGB = {
    StatusConformidade.ATENDIDO: (11, 107, 58),       # #0b6b3a
    StatusConformidade.PARCIAL: (95, 125, 18),        # #5f7d12
    StatusConformidade.NAO_ATENDIDO: (155, 28, 28),   # #9b1c1c
}

STATUS_HEX = {
    StatusConformidade.ATENDIDO: "0B6B3A",
    StatusConformidade.PARCIAL: "5F7D12",
    StatusConformidade.NAO_ATENDIDO: "9B1C1C",
}

_LOGO_PATH = Path(__file__).resolve().parent.parent / "app" / "static" / "logo_codevasf.png"


def _meta_rows(relatorio: RelatorioConformidade) -> list[tuple[str, str]]:
    counts = relatorio.contagem
    agora = datetime.now().strftime("%d/%m/%Y %H:%M")
    return [
        ("Data", agora),
        ("Tipo", label_tipo(relatorio.tipo)),
        ("Entidade / Município", relatorio.entidade_detectada),
        ("Resumo", relatorio.resumo),
        ("Atendidos", str(counts["atendido"])),
        ("Parciais", str(counts["parcial"])),
        ("Não Atendidos", str(counts["nao_atendido"])),
        ("Versão", "Revisada (Humano)" if relatorio.revisado else "Automática"),
    ]


def _find_unicode_font() -> str | None:
    candidates = [
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
        "/usr/share/fonts/TTF/DejaVuSans.ttf",
        str(Path.home() / "miniconda3/fonts/DejaVuSans.ttf"),
    ]
    for path in candidates:
        if Path(path).is_file():
            return path
    return None


def _thin_border() -> Border:
    return Border(
        left=Side(style="thin", color=BORDA),
        right=Side(style="thin", color=BORDA),
        top=Side(style="thin", color=BORDA),
        bottom=Side(style="thin", color=BORDA),
    )


def relatorio_para_markdown(relatorio: RelatorioConformidade) -> str:
    counts = relatorio.contagem
    lines = [
        "# Codevasf 12ª SR — Relatório de Conformidade Documental",
        "",
        "## Resumo",
        "",
        f"| Atendidos | Parciais | Não Atendidos |",
        f"| :---: | :---: | :---: |",
        f"| **{counts['atendido']}** | **{counts['parcial']}** | **{counts['nao_atendido']}** |",
        "",
        "### Legenda",
        "",
        "- **Atendido** — Documento Ok",
        "- **Parcial** — Incompleto / Dúvida",
        "- **Não Atendido** — Ausente",
        "",
    ]
    for label, value in _meta_rows(relatorio):
        if label in {"Atendidos", "Parciais", "Não Atendidos"}:
            continue
        lines.append(f"- **{label}:** {value}")

    lines.extend(["", "## Itens do Checklist", ""])
    for item in relatorio.itens:
        lines.append(
            f"### {item.numero}. [{STATUS_LABEL[item.status]}] "
            f"({item.fonte}) {item.descricao}"
        )
        lines.append("")
        lines.append(f"**Motivo:** {item.motivo}")
        if item.documentos_relacionados:
            lines.append("")
            lines.append("**Arquivos:** " + ", ".join(item.documentos_relacionados))
        lines.append("")

    lines.extend(["## Documentos Analisados", ""])
    for name in relatorio.documentos_analisados:
        lines.append(f"- `{name}`")

    lines.extend(
        [
            "",
            "---",
            "_Relatório assistivo — Codevasf 12ª SR (Natal/RN). "
            "A decisão final permanece com a equipe técnica._",
        ]
    )
    return "\n".join(lines)


def relatorio_para_xlsx(relatorio: RelatorioConformidade) -> bytes:
    """Planilha com cards de resumo e status coloridos (como no frontend)."""
    counts = relatorio.contagem
    agora = datetime.now().strftime("%d/%m/%Y %H:%M")
    thin = _thin_border()
    wrap = Alignment(wrap_text=True, vertical="top")
    center = Alignment(horizontal="center", vertical="center", wrap_text=True)

    wb = Workbook()
    ws = wb.active
    ws.title = "Resumo"

    # Cabeçalho institucional
    ws.merge_cells("A1:F1")
    ws["A1"] = "Codevasf 12ª SR — Relatório de Conformidade Documental"
    ws["A1"].font = Font(bold=True, size=16, color="FFFFFF")
    ws["A1"].fill = PatternFill("solid", fgColor=AZUL_MARCA)
    ws["A1"].alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[1].height = 28

    ws.merge_cells("A2:F2")
    ws["A2"] = (
        f"{label_tipo(relatorio.tipo)} · {relatorio.entidade_detectada} · "
        f"{'Versão Revisada' if relatorio.revisado else 'Versão Automática'} · {agora}"
    )
    ws["A2"].font = Font(size=10, color=AZUL_ESCURO)
    ws["A2"].fill = PatternFill("solid", fgColor=OFFWHITE)
    ws["A2"].alignment = Alignment(horizontal="center", vertical="center")

    # Cards de contagem (como no frontend)
    ws["A4"] = "Atendidos"
    ws["C4"] = "Parciais"
    ws["E4"] = "Não Atendidos"
    ws["A5"] = counts["atendido"]
    ws["C5"] = counts["parcial"]
    ws["E5"] = counts["nao_atendido"]
    for col, status in (
        ("A", StatusConformidade.ATENDIDO),
        ("C", StatusConformidade.PARCIAL),
        ("E", StatusConformidade.NAO_ATENDIDO),
    ):
        for row in (4, 5):
            cell = ws[f"{col}{row}"]
            cell.fill = CARD_FILL[status]
            cell.alignment = center
            cell.border = thin
            cell.font = Font(
                bold=True,
                size=16 if row == 5 else 10,
                color=STATUS_HEX[status],
            )
        ws.merge_cells(f"{col}4:{chr(ord(col)+1)}4")
        ws.merge_cells(f"{col}5:{chr(ord(col)+1)}5")
        ws[f"{chr(ord(col)+1)}4"].border = thin
        ws[f"{chr(ord(col)+1)}5"].border = thin
        ws[f"{chr(ord(col)+1)}4"].fill = CARD_FILL[status]
        ws[f"{chr(ord(col)+1)}5"].fill = CARD_FILL[status]
    ws.row_dimensions[5].height = 32

    # Legenda
    ws["A7"] = "Legenda"
    ws["A7"].font = Font(bold=True, color=AZUL_MARCA)
    ws["A8"] = "Atendido — Documento Ok"
    ws["A8"].fill = STATUS_FILL[StatusConformidade.ATENDIDO]
    ws["C8"] = "Parcial — Incompleto / Dúvida"
    ws["C8"].fill = STATUS_FILL[StatusConformidade.PARCIAL]
    ws["E8"] = "Não Atendido — Ausente"
    ws["E8"].fill = STATUS_FILL[StatusConformidade.NAO_ATENDIDO]
    legend_fills = {
        "A": STATUS_FILL[StatusConformidade.ATENDIDO],
        "C": STATUS_FILL[StatusConformidade.PARCIAL],
        "E": STATUS_FILL[StatusConformidade.NAO_ATENDIDO],
    }
    for col in ("A", "C", "E"):
        ws.merge_cells(f"{col}8:{chr(ord(col)+1)}8")
        ws[f"{col}8"].border = thin
        ws[f"{col}8"].alignment = center
        ws[f"{chr(ord(col)+1)}8"].fill = legend_fills[col]
        ws[f"{chr(ord(col)+1)}8"].border = thin

    ws["A10"] = "Resumo narrativo"
    ws["A10"].font = Font(bold=True, color=AZUL_ESCURO)
    ws.merge_cells("A11:F11")
    ws["A11"] = relatorio.resumo
    ws["A11"].alignment = wrap
    ws["A11"].border = thin
    ws.row_dimensions[11].height = 48

    # Aba Itens
    ws_itens = wb.create_sheet("Itens")
    headers = [
        "Nº",
        "Status",
        "Legenda",
        "Fonte",
        "Descrição do Item",
        "Motivo",
        "Arquivos Relacionados",
    ]
    header_fill = PatternFill("solid", fgColor=AZUL_MARCA)
    header_font = Font(bold=True, color="FFFFFF")
    for col, title in enumerate(headers, start=1):
        cell = ws_itens.cell(row=1, column=col, value=title)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(vertical="center", horizontal="center")
        cell.border = thin

    for row_idx, item in enumerate(relatorio.itens, start=2):
        values = [
            item.numero,
            STATUS_LABEL[item.status],
            STATUS_LEGENDA[item.status],
            item.fonte.title() if item.fonte else "",
            item.descricao,
            item.motivo,
            ", ".join(item.documentos_relacionados),
        ]
        for col, value in enumerate(values, start=1):
            cell = ws_itens.cell(row=row_idx, column=col, value=value)
            cell.border = thin
            cell.alignment = wrap
            if col in (2, 3):
                cell.fill = STATUS_FILL[item.status]
                cell.font = Font(bold=True, color=STATUS_HEX[item.status])
                cell.alignment = center
        ws_itens.row_dimensions[row_idx].height = 48

    widths = [6, 14, 22, 10, 42, 42, 32]
    for idx, width in enumerate(widths, start=1):
        ws_itens.column_dimensions[get_column_letter(idx)].width = width
    ws_itens.auto_filter.ref = f"A1:G{max(1, len(relatorio.itens) + 1)}"
    ws_itens.freeze_panes = "A2"

    # Aba Documentos
    ws_docs = wb.create_sheet("Documentos")
    cell = ws_docs.cell(row=1, column=1, value="Arquivo Analisado")
    cell.font = header_font
    cell.fill = header_fill
    cell.border = thin
    for row_idx, name in enumerate(relatorio.documentos_analisados, start=2):
        cell = ws_docs.cell(row=row_idx, column=1, value=name)
        cell.border = thin
        cell.fill = PatternFill("solid", fgColor=OFFWHITE)
    ws_docs.column_dimensions["A"].width = 70

    for col in "ABCDEF":
        ws.column_dimensions[col].width = 14
    ws.column_dimensions["B"].width = 16
    ws.column_dimensions["D"].width = 16
    ws.column_dimensions["F"].width = 16

    ws["A13"] = (
        "Codevasf — 12ª Superintendência Regional (Natal/RN). "
        "Relatório assistivo; a decisão final permanece com a equipe técnica."
    )
    ws["A13"].font = Font(italic=True, color="666666", size=9)
    ws.merge_cells("A13:F13")

    buffer = BytesIO()
    wb.save(buffer)
    return buffer.getvalue()


def relatorio_para_docx(relatorio: RelatorioConformidade) -> bytes:
    """Word com cabeçalho institucional, cards de resumo e status coloridos."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import Inches, Pt, RGBColor

    def _set_cell_shading(cell, hex_color: str) -> None:
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), hex_color)
        shading.set(qn("w:val"), "clear")
        cell._tc.get_or_add_tcPr().append(shading)

    def _rgb(hex6: str) -> RGBColor:
        return RGBColor(int(hex6[0:2], 16), int(hex6[2:4], 16), int(hex6[4:6], 16))

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)

    if _LOGO_PATH.is_file():
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_logo.add_run().add_picture(str(_LOGO_PATH), width=Inches(2.4))

    title = doc.add_heading(
        "Codevasf 12ª SR — Relatório de Conformidade Documental", level=1
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = _rgb(AZUL_MARCA)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_s = subtitle.add_run(
        f"{label_tipo(relatorio.tipo)} · {relatorio.entidade_detectada} · "
        f"{'Versão Revisada' if relatorio.revisado else 'Versão Automática'}"
    )
    run_s.font.size = Pt(10)
    run_s.font.color.rgb = _rgb(AZUL_ESCURO)

    # Tabela-resumo (3 cards)
    counts = relatorio.contagem
    table = doc.add_table(rows=2, cols=3)
    table.style = "Table Grid"
    headers = ("Atendidos", "Parciais", "Não Atendidos")
    values = (counts["atendido"], counts["parcial"], counts["nao_atendido"])
    fills = ("EAF8F0", "F7F9E8", "FDECED")
    colors = (
        STATUS_HEX[StatusConformidade.ATENDIDO],
        STATUS_HEX[StatusConformidade.PARCIAL],
        STATUS_HEX[StatusConformidade.NAO_ATENDIDO],
    )
    for i in range(3):
        cell_h = table.rows[0].cells[i]
        cell_v = table.rows[1].cells[i]
        cell_h.text = headers[i]
        cell_v.text = str(values[i])
        _set_cell_shading(cell_h, fills[i])
        _set_cell_shading(cell_v, fills[i])
        for cell in (cell_h, cell_v):
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.bold = True
                    run.font.color.rgb = _rgb(colors[i])
                    if cell is cell_v:
                        run.font.size = Pt(18)

    doc.add_paragraph()
    leg = doc.add_paragraph()
    r = leg.add_run("Legenda: ")
    r.bold = True
    r.font.color.rgb = _rgb(AZUL_MARCA)
    leg.add_run(
        "Atendido — Documento Ok · Parcial — Incompleto / Dúvida · "
        "Não Atendido — Ausente"
    )

    for label, value in _meta_rows(relatorio):
        if label in {"Atendidos", "Parciais", "Não Atendidos", "Tipo", "Entidade / Município", "Versão"}:
            continue
        p = doc.add_paragraph()
        run_l = p.add_run(f"{label}: ")
        run_l.bold = True
        run_l.font.color.rgb = _rgb(AZUL_ESCURO)
        p.add_run(value)

    doc.add_heading("Itens do Checklist", level=2)
    badge_bg = {
        StatusConformidade.ATENDIDO: "D7F0E5",
        StatusConformidade.PARCIAL: "EEF7D4",
        StatusConformidade.NAO_ATENDIDO: "FDE2E2",
    }
    for item in relatorio.itens:
        item_table = doc.add_table(rows=1, cols=1)
        item_table.style = "Table Grid"
        cell = item_table.rows[0].cells[0]
        _set_cell_shading(cell, badge_bg[item.status])
        cell.paragraphs[0].clear()
        p_status = cell.paragraphs[0]
        run_n = p_status.add_run(
            f"{item.numero}. [{STATUS_LABEL[item.status]}] ({item.fonte}) "
        )
        run_n.bold = True
        run_n.font.color.rgb = RGBColor(*STATUS_RGB[item.status])
        run_rest = p_status.add_run(item.descricao)
        run_rest.font.color.rgb = _rgb(AZUL_ESCURO)

        p_motivo = doc.add_paragraph()
        r = p_motivo.add_run("Motivo: ")
        r.bold = True
        p_motivo.add_run(item.motivo)
        if item.documentos_relacionados:
            p_arq = doc.add_paragraph()
            r2 = p_arq.add_run("Arquivos: ")
            r2.bold = True
            p_arq.add_run(", ".join(item.documentos_relacionados))
        doc.add_paragraph()

    doc.add_heading("Documentos Analisados", level=2)
    for name in relatorio.documentos_analisados:
        doc.add_paragraph(name, style="List Bullet")

    footnote = doc.add_paragraph()
    run_f = footnote.add_run(
        "Codevasf — 12ª Superintendência Regional (Natal/RN). "
        "Relatório assistivo; a decisão final permanece com a equipe técnica."
    )
    run_f.italic = True
    run_f.font.size = Pt(9)
    run_f.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def relatorio_para_pdf(relatorio: RelatorioConformidade) -> bytes:
    """PDF com faixa institucional, cards de resumo e status coloridos."""
    from fpdf import FPDF

    class RelatorioPDF(FPDF):
        def header(self) -> None:
            # Faixa azul institucional
            self.set_fill_color(0, 92, 168)
            self.rect(0, 0, self.w, 16, "F")
            self.set_xy(10, 5)
            self.set_text_color(255, 255, 255)
            self.set_font(self._body_font, "B", 11)
            self.cell(0, 6, "Codevasf 12ª SR — Conformidade Documental", align="C")
            self.ln(14)
            self.set_text_color(0, 0, 0)

        def footer(self) -> None:
            self.set_y(-15)
            self.set_font(self._body_font, size=8)
            self.set_text_color(100, 100, 100)
            self.cell(
                0,
                8,
                "Codevasf 12ª SR (Natal/RN) — relatório assistivo — decisão final da equipe técnica",
                align="C",
            )

    pdf = RelatorioPDF()
    pdf.set_auto_page_break(auto=True, margin=18)
    font_path = _find_unicode_font()
    if font_path:
        pdf.add_font("Body", "", font_path)
        bold_path = font_path.replace("DejaVuSans.ttf", "DejaVuSans-Bold.ttf")
        if Path(bold_path).is_file():
            pdf.add_font("Body", "B", bold_path)
        else:
            pdf.add_font("Body", "B", font_path)
        pdf._body_font = "Body"  # type: ignore[attr-defined]
    else:
        pdf._body_font = "Helvetica"  # type: ignore[attr-defined]

    body = pdf._body_font  # type: ignore[attr-defined]
    pdf.add_page()

    def _reset_x() -> None:
        pdf.set_x(pdf.l_margin)

    if _LOGO_PATH.is_file():
        try:
            logo_w = 50.0
            pdf.image(str(_LOGO_PATH), x=(pdf.w - logo_w) / 2, w=logo_w)
            pdf.ln(6)
        except Exception:
            pass

    _reset_x()
    pdf.set_font(body, "B", 13)
    pdf.set_text_color(0, 92, 168)
    pdf.multi_cell(0, 7, "Relatório de Conformidade Documental", align="C")
    _reset_x()
    pdf.set_font(body, "", 9)
    pdf.set_text_color(34, 43, 84)
    pdf.multi_cell(
        0,
        5,
        f"{label_tipo(relatorio.tipo)} · {relatorio.entidade_detectada} · "
        f"{'Versão Revisada' if relatorio.revisado else 'Versão Automática'}",
        align="C",
    )
    pdf.ln(3)

    # Cards de resumo (mesmas cores do frontend)
    counts = relatorio.contagem
    usable = pdf.w - pdf.l_margin - pdf.r_margin
    gap = 4.0
    card_w = (usable - 2 * gap) / 3
    card_h = 18.0
    x0 = pdf.l_margin
    y0 = pdf.get_y()
    cards = [
        ("Atendidos", counts["atendido"], (234, 248, 240), STATUS_RGB[StatusConformidade.ATENDIDO]),
        ("Parciais", counts["parcial"], (247, 249, 232), STATUS_RGB[StatusConformidade.PARCIAL]),
        ("Não Atendidos", counts["nao_atendido"], (253, 236, 237), STATUS_RGB[StatusConformidade.NAO_ATENDIDO]),
    ]
    for i, (label, value, bg, fg) in enumerate(cards):
        x = x0 + i * (card_w + gap)
        pdf.set_fill_color(*bg)
        pdf.set_draw_color(183, 215, 232)
        pdf.rect(x, y0, card_w, card_h, "FD")
        pdf.set_xy(x, y0 + 2)
        pdf.set_font(body, "B", 9)
        pdf.set_text_color(*fg)
        pdf.cell(card_w, 5, label, align="C")
        pdf.set_xy(x, y0 + 8)
        pdf.set_font(body, "B", 16)
        pdf.cell(card_w, 8, str(value), align="C")
    pdf.set_xy(pdf.l_margin, y0 + card_h + 4)

    pdf.set_font(body, "", 8)
    pdf.set_text_color(74, 88, 120)
    pdf.multi_cell(
        0,
        4,
        "Legenda: Atendido — Documento Ok · Parcial — Incompleto / Dúvida · "
        "Não Atendido — Ausente",
        align="C",
    )
    pdf.ln(2)

    pdf.set_text_color(34, 43, 84)
    for label, value in _meta_rows(relatorio):
        if label in {"Atendidos", "Parciais", "Não Atendidos", "Tipo", "Entidade / Município", "Versão"}:
            continue
        _reset_x()
        pdf.set_font(body, "B", 9)
        pdf.multi_cell(0, 5, f"{label}: {value}")

    pdf.ln(2)
    _reset_x()
    pdf.set_font(body, "B", 12)
    pdf.set_text_color(0, 92, 168)
    pdf.multi_cell(0, 7, "Itens do Checklist")
    pdf.ln(1)

    for item in relatorio.itens:
        r, g, b = STATUS_RGB[item.status]
        _reset_x()
        pdf.set_fill_color(r, g, b)
        pdf.set_text_color(255, 255, 255)
        pdf.set_font(body, "B", 9)
        pdf.cell(
            0,
            6,
            f"  {item.numero}. {STATUS_LABEL[item.status]}  ({item.fonte})",
            new_x="LMARGIN",
            new_y="NEXT",
            fill=True,
        )
        pdf.set_text_color(34, 43, 84)
        pdf.set_font(body, "", 9)
        _reset_x()
        pdf.multi_cell(0, 5, item.descricao)
        _reset_x()
        pdf.set_font(body, "", 8)
        pdf.multi_cell(0, 4, f"Motivo: {item.motivo}")
        if item.documentos_relacionados:
            _reset_x()
            pdf.multi_cell(0, 4, "Arquivos: " + ", ".join(item.documentos_relacionados))
        pdf.ln(2)

    _reset_x()
    pdf.set_font(body, "B", 12)
    pdf.set_text_color(0, 92, 168)
    pdf.multi_cell(0, 7, "Documentos Analisados")
    pdf.set_text_color(34, 43, 84)
    pdf.set_font(body, "", 9)
    for name in relatorio.documentos_analisados:
        _reset_x()
        pdf.multi_cell(0, 4, f"• {name}")

    return bytes(pdf.output())
