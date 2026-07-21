"""
Carregamento e extração de texto de documentos do requerimento.

Fontes suportadas: ZIP, pasta ou arquivos avulsos (PDF, DOCX, TXT, imagens).
PDFs com pouco texto nativo e imagens passam por OCR (ver ``conformidade.ocr``).

Funções-chave: ``load_from_zip``, ``scan_folder``, ``ocr_available``.
"""

from __future__ import annotations

import shutil
import zipfile
from dataclasses import dataclass
from pathlib import Path

from conformidade.ocr import (
    ocr_available,
    ocr_image_file,
    ocr_pdf,
    pdf_needs_any_ocr,
)

SUPPORTED_TEXT_EXTENSIONS = {".pdf", ".docx", ".txt", ".md", ".doc"}
SUPPORTED_IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp", ".webp"}
SUPPORTED_EXTENSIONS = SUPPORTED_TEXT_EXTENSIONS | SUPPORTED_IMAGE_EXTENSIONS

# Reexport para UI / scripts que importam de loaders
__all__ = [
    "LoadedDocument",
    "SUPPORTED_EXTENSIONS",
    "extract_zip",
    "load_file",
    "load_from_zip",
    "ocr_available",
    "save_uploaded_bytes",
    "scan_folder",
    "summarize_inventory",
]


@dataclass
class LoadedDocument:
    source: str
    content: str
    file_name: str
    relative_path: str
    extraction_method: str = "texto"  # texto | ocr | hibrido | vazio | erro


def _read_text_file(path: Path) -> str:
    for encoding in ("utf-8", "latin-1", "cp1252"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    return path.read_text(encoding="utf-8", errors="replace")


def _read_pdf_native(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    pages: list[str] = []
    for page in reader.pages:
        text = page.extract_text() or ""
        if text.strip():
            pages.append(text)
    return "\n\n".join(pages)


def _read_pdf(path: Path) -> tuple[str, str]:
    """Retorna (conteúdo, método). Método: texto | ocr | hibrido | vazio | erro."""
    native = ""
    page_count = None
    try:
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        page_count = len(reader.pages)
        native = _read_pdf_native(path)
    except Exception:
        native = ""

    needs_ocr = pdf_needs_any_ocr(native, page_count)
    # Mesmo com texto nativo “suficiente” no agregado, ocr_pdf só OCR nas páginas pobres.
    # Se o agregado parece bom, ainda assim valida página a página via ocr_pdf quando
    # há indício de scan (texto curto ou vazio).
    if not needs_ocr:
        # Checagem rápida por página (fitz): se alguma página for pobre, usa pipeline híbrido
        try:
            from conformidade.ocr import page_native_texts, page_needs_ocr

            pages = page_native_texts(path)
            if any(page_needs_ocr(p) for p in pages):
                needs_ocr = True
            else:
                return native.strip(), "texto"
        except Exception:
            return native.strip(), "texto"

    ok, _msg = ocr_available()
    if not ok:
        if native.strip():
            return native.strip(), "texto"
        return (
            f"[Arquivo sem texto extraível e OCR indisponível: {path.name}. "
            "Instale: apt install tesseract-ocr tesseract-ocr-por tesseract-ocr-eng && "
            "pip install pymupdf pytesseract pillow]",
            "vazio",
        )

    try:
        content, method = ocr_pdf(path)
    except Exception as exc:
        if native.strip():
            return native.strip(), "texto"
        return f"[Falha no OCR de {path.name}: {exc}]", "erro"

    content_s = (content or "").strip()
    native_s = native.strip()
    if content_s:
        return content_s, method
    if native_s:
        return native_s, "texto"
    return f"[OCR não extraiu texto de {path.name}.]", "vazio"


def _read_docx(path: Path) -> str:
    from docx import Document

    document = Document(str(path))
    paragraphs = [
        paragraph.text.strip()
        for paragraph in document.paragraphs
        if paragraph.text.strip()
    ]
    return "\n\n".join(paragraphs)


def load_file(path: Path, root: Path | None = None) -> LoadedDocument | None:
    suffix = path.suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        return None

    method = "texto"
    try:
        if suffix == ".pdf":
            content, method = _read_pdf(path)
        elif suffix in SUPPORTED_IMAGE_EXTENSIONS:
            ok, msg = ocr_available()
            if not ok:
                content = f"[Imagem {path.name}: OCR indisponível — {msg}]"
                method = "vazio"
            else:
                content = ocr_image_file(path)
                method = "ocr" if content.strip() else "vazio"
                if not content.strip():
                    content = f"[OCR não extraiu texto da imagem {path.name}.]"
        elif suffix in {".docx", ".doc"}:
            if suffix == ".doc":
                content = (
                    f"[Arquivo .doc binário: {path.name}. "
                    "Conteúdo não extraído automaticamente.]"
                )
                method = "vazio"
            else:
                content = _read_docx(path)
        else:
            content = _read_text_file(path)
    except Exception as exc:
        content = f"[Falha ao ler {path.name}: {exc}]"
        method = "erro"

    content = (content or "").strip()
    if not content:
        content = f"[Arquivo sem texto extraível: {path.name}.]"
        method = "vazio"

    rel = str(path.relative_to(root)) if root else path.name
    return LoadedDocument(
        source=str(path),
        content=content,
        file_name=path.name,
        relative_path=rel,
        extraction_method=method,
    )


def scan_folder(root: Path) -> list[LoadedDocument]:
    if not root.exists():
        raise FileNotFoundError(f"Pasta não encontrada: {root}")
    if not root.is_dir():
        raise NotADirectoryError(f"Caminho não é uma pasta: {root}")

    documents: list[LoadedDocument] = []
    for path in sorted(root.rglob("*")):
        if not path.is_file():
            continue
        if path.name.startswith(".") or path.name.startswith("__MACOSX"):
            continue
        loaded = load_file(path, root=root)
        if loaded is not None:
            documents.append(loaded)
    return documents


def extract_zip(zip_path: Path, destination: Path) -> Path:
    """Extrai ZIP para destination e retorna a pasta raiz útil."""
    destination.mkdir(parents=True, exist_ok=True)
    extract_dir = destination / zip_path.stem
    if extract_dir.exists():
        shutil.rmtree(extract_dir)
    extract_dir.mkdir(parents=True, exist_ok=True)

    with zipfile.ZipFile(zip_path, "r") as archive:
        archive.extractall(extract_dir)

    children = [p for p in extract_dir.iterdir() if not p.name.startswith("__MACOSX")]
    if len(children) == 1 and children[0].is_dir():
        return children[0]
    return extract_dir


def load_from_zip(zip_path: Path, work_dir: Path) -> list[LoadedDocument]:
    if not zip_path.exists():
        raise FileNotFoundError(f"ZIP não encontrado: {zip_path}")
    root = extract_zip(zip_path, work_dir)
    return scan_folder(root)


def save_uploaded_bytes(file_name: str, data: bytes, destination: Path) -> Path:
    destination.mkdir(parents=True, exist_ok=True)
    target = destination / file_name
    target.write_bytes(data)
    return target


def summarize_inventory(documents: list[LoadedDocument]) -> str:
    if not documents:
        return "(nenhum documento encontrado)"
    lines = []
    for doc in documents:
        lines.append(
            f"- {doc.relative_path} ({len(doc.content)} caracteres, método: {doc.extraction_method})"
        )
    ocr_count = sum(1 for d in documents if d.extraction_method in {"ocr", "hibrido"})
    if ocr_count:
        lines.append(f"\n({ocr_count} arquivo(s) lido(s) via OCR/híbrido)")
    return "\n".join(lines)
